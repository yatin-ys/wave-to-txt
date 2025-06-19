import os
import tempfile
from typing import Optional
from langchain.schema import Document
import filetype
from pypdf import PdfReader
from docx import Document as DocxDocument

from backend.core.logging_config import get_logger

logger = get_logger("document_processor")

SUPPORTED_FILE_TYPES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "text/plain": ".txt",
}


class LightweightDocumentProcessor:
    """
    Lightweight document processor for PDF, DOCX, and TXT files.
    Uses basic libraries instead of heavy unstructured library.
    """

    @staticmethod
    def validate_file_type(file_path: str) -> bool:
        """
        Validate if the file type is supported.

        Args:
            file_path: Path to the file to validate

        Returns:
            bool: True if file type is supported, False otherwise
        """
        try:
            kind = filetype.guess(file_path)
            if kind is None:
                # Try to determine from extension for text files
                if file_path.lower().endswith(".txt"):
                    return True
                logger.warning(f"Could not determine file type: {file_path}")
                return False

            mime_type = kind.mime
            is_supported = mime_type in SUPPORTED_FILE_TYPES

            logger.debug(
                f"File type validation",
                extra={
                    "file_path": file_path,
                    "mime_type": mime_type,
                    "is_supported": is_supported,
                },
            )

            return is_supported

        except Exception as e:
            logger.error(
                f"Error validating file type: {e}",
                exc_info=True,
                extra={"file_path": file_path},
            )
            return False

    @staticmethod
    def process_pdf(file_path: str) -> Document:
        """
        Extract text from PDF file using pypdf.

        Args:
            file_path: Path to the PDF file

        Returns:
            Document: LangChain document with extracted text and metadata
        """
        try:
            reader = PdfReader(file_path)
            text_content = []

            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(f"[Page {page_num + 1}]\n{page_text}")

            full_text = "\n\n".join(text_content)

            if not full_text.strip():
                raise ValueError("No text content could be extracted from the PDF")

            metadata = {
                "file_type": "pdf",
                "pages": len(reader.pages),
                "source": os.path.basename(file_path),
            }

            logger.info(
                f"Successfully processed PDF",
                extra={
                    "file_path": file_path,
                    "pages": len(reader.pages),
                    "text_length": len(full_text),
                },
            )

            return Document(page_content=full_text, metadata=metadata)

        except Exception as e:
            logger.error(
                f"Error processing PDF: {e}",
                exc_info=True,
                extra={"file_path": file_path},
            )
            raise

    @staticmethod
    def process_docx(file_path: str) -> Document:
        """
        Extract text from DOCX file using python-docx.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Document: LangChain document with extracted text and metadata
        """
        try:
            doc = DocxDocument(file_path)
            paragraphs = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)

            full_text = "\n\n".join(paragraphs)

            if not full_text.strip():
                raise ValueError("No text content could be extracted from the DOCX")

            metadata = {
                "file_type": "docx",
                "paragraphs": len(paragraphs),
                "source": os.path.basename(file_path),
            }

            logger.info(
                f"Successfully processed DOCX",
                extra={
                    "file_path": file_path,
                    "paragraphs": len(paragraphs),
                    "text_length": len(full_text),
                },
            )

            return Document(page_content=full_text, metadata=metadata)

        except Exception as e:
            logger.error(
                f"Error processing DOCX: {e}",
                exc_info=True,
                extra={"file_path": file_path},
            )
            raise

    @staticmethod
    def process_txt(file_path: str) -> Document:
        """
        Extract text from TXT file.

        Args:
            file_path: Path to the TXT file

        Returns:
            Document: LangChain document with extracted text and metadata
        """
        try:
            encodings = ["utf-8", "latin-1", "cp1252"]
            text_content = None

            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        text_content = file.read()
                    break
                except UnicodeDecodeError:
                    continue

            if text_content is None:
                raise ValueError(
                    "Could not decode the text file with any supported encoding"
                )

            if not text_content.strip():
                raise ValueError(
                    "The text file is empty or contains no readable content"
                )

            metadata = {"file_type": "txt", "source": os.path.basename(file_path)}

            logger.info(
                f"Successfully processed TXT",
                extra={"file_path": file_path, "text_length": len(text_content)},
            )

            return Document(page_content=text_content, metadata=metadata)

        except Exception as e:
            logger.error(
                f"Error processing TXT: {e}",
                exc_info=True,
                extra={"file_path": file_path},
            )
            raise

    @classmethod
    def process_file(cls, file_path: str) -> Optional[Document]:
        """
        Process a file and extract its content based on file type.

        Args:
            file_path: Path to the file to process

        Returns:
            Optional[Document]: LangChain document with extracted content,
                               or None if processing failed
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")

            if not cls.validate_file_type(file_path):
                raise ValueError(f"Unsupported file type: {file_path}")

            # Determine file type and process accordingly
            kind = filetype.guess(file_path)

            if kind and kind.mime == "application/pdf":
                return cls.process_pdf(file_path)
            elif (
                kind
                and kind.mime
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                return cls.process_docx(file_path)
            elif kind is None and file_path.lower().endswith(".txt"):
                return cls.process_txt(file_path)
            else:
                raise ValueError(
                    f"Unsupported file type: {kind.mime if kind else 'unknown'}"
                )

        except Exception as e:
            logger.error(
                f"Failed to process file: {e}",
                exc_info=True,
                extra={"file_path": file_path},
            )
            return None

    @classmethod
    def process_uploaded_file(
        cls, file_content: bytes, filename: str
    ) -> Optional[Document]:
        """
        Process an uploaded file from bytes content.

        Args:
            file_content: Raw bytes content of the uploaded file
            filename: Original filename

        Returns:
            Optional[Document]: LangChain document with extracted content
        """
        temp_file = None
        try:
            # Create temporary file
            suffix = os.path.splitext(filename)[1]
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            # Process the temporary file
            document = cls.process_file(temp_file_path)

            if document:
                # Update metadata with original filename
                document.metadata["original_filename"] = filename

            return document

        except Exception as e:
            logger.error(
                f"Failed to process uploaded file: {e}",
                exc_info=True,
                extra={"filename": filename},
            )
            return None

        finally:
            # Clean up temporary file
            if temp_file and os.path.exists(temp_file.name):
                try:
                    os.unlink(temp_file.name)
                except Exception as e:
                    logger.warning(f"Failed to clean up temporary file: {e}")

    @staticmethod
    def get_supported_extensions() -> list[str]:
        """
        Get list of supported file extensions.

        Returns:
            List[str]: List of supported file extensions
        """
        return [".pdf", ".docx", ".txt"]

    @staticmethod
    def get_supported_mime_types() -> list[str]:
        """
        Get list of supported MIME types.

        Returns:
            List[str]: List of supported MIME types
        """
        return list(SUPPORTED_FILE_TYPES.keys())


# Global instance
document_processor = LightweightDocumentProcessor()
